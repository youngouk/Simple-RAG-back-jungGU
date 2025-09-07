"""
Document processing module
문서 로딩, 분할, 임베딩 처리 모듈
"""
import asyncio
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
import mimetypes
import hashlib

# Document loaders
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup
import markdown
import json

# LangChain components
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from fastembed import SparseTextEmbedding

from ..lib.logger import get_logger

logger = get_logger(__name__)

# LangChain JSONLoader (선택적 import)
try:
    from langchain_community.document_loaders import JSONLoader
    LANGCHAIN_JSONLOADER_AVAILABLE = True
except ImportError:
    logger.warning("LangChain JSONLoader not available, using basic JSON processing")
    LANGCHAIN_JSONLOADER_AVAILABLE = False

class DocumentProcessor:
    """문서 처리 모듈"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.document_config = config.get('document_processing', {})
        self.embeddings_config = config.get('embeddings', {})
        
        # 지원하는 파일 타입
        self.supported_types = self.document_config.get('file_types', [
            'pdf', 'txt', 'docx', 'xlsx', 'csv', 'html', 'md', 'json'
        ])
        
        # 텍스트 분할 설정
        self.chunk_size = self.document_config.get('chunk_size', 400)
        self.chunk_overlap = self.document_config.get('chunk_overlap', 50)
        
        # 임베딩 모델 초기화
        self.embedder = None
        self.sparse_embedder = None
        self._init_embedders()
        
    def _init_embedders(self):
        """Dense와 Sparse 임베딩 모델 초기화"""
        try:
            # Dense embeddings
            provider = self.embeddings_config.get('provider', 'google')
            
            if provider == 'gemini':
                # Gemini Embedding 001 사용 (1536차원)
                from .gemini_embeddings import GeminiEmbeddings
                model_name = self.embeddings_config.get('model', 'models/gemini-embedding-001')
                output_dimensionality = self.embeddings_config.get('output_dimensionality', 1536)
                batch_size = self.embeddings_config.get('batch_size', 100)
                # config.yaml의 llm.google.api_key에서 가져오기
                api_key = self.config.get('llm', {}).get('google', {}).get('api_key')
                
                self.embedder = GeminiEmbeddings(
                    google_api_key=api_key,
                    model_name=model_name,
                    output_dimensionality=output_dimensionality,
                    batch_size=batch_size
                )
                logger.info(f"Gemini embeddings initialized with model: {model_name}, dimensions: {output_dimensionality}")
                
            elif provider == 'google':
                # 기존 Google text-embedding-004 지원 유지 (호환성)
                model_name = self.embeddings_config.get('model', 'text-embedding-004')
                # config.yaml의 llm.google.api_key에서 가져오기
                api_key = self.config.get('llm', {}).get('google', {}).get('api_key')
                self.embedder = GoogleGenerativeAIEmbeddings(
                    model=model_name,
                    google_api_key=api_key
                )
                logger.info(f"Google embeddings initialized with model: {model_name}")
            else:
                raise ValueError(f"Unsupported embedding provider: {provider}")
            
            # Sparse embeddings (BM42)
            sparse_model = self.embeddings_config.get('sparse_model', 'Qdrant/bm42-all-minilm-l6-v2-attentions')
            self.sparse_embedder = SparseTextEmbedding(model_name=sparse_model)
            logger.info(f"Sparse embeddings initialized with model: {sparse_model}")
                
        except Exception as e:
            logger.error(f"Failed to initialize embedders: {e}")
            raise
    
    async def load_document(self, file_path: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """문서 로드"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # 파일 타입 결정
        file_type = self._get_file_type(file_path)
        
        if file_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        logger.info(f"Loading document: {file_path.name} (type: {file_type})")
        
        try:
            # 파일 타입별 로더 호출
            documents = await self._load_by_type(file_path, file_type)
            
            # 파일 크기 계산
            file_size = file_path.stat().st_size
            
            # 메타데이터 추가
            for i, doc in enumerate(documents):
                doc.metadata.update({
                    'source_file': file_path.name,
                    'file_type': file_type,
                    'file_path': str(file_path),
                    'file_size': file_size,
                    'chunk_index': i,
                    'total_chunks': len(documents),
                    'file_hash': self._get_file_hash(file_path),
                    'load_timestamp': asyncio.get_event_loop().time(),
                    **(metadata or {})
                })
            
            logger.info(f"Document loaded successfully: {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise
    
    def _get_file_type(self, file_path: Path) -> str:
        """파일 타입 결정"""
        # 확장자 기반
        ext = file_path.suffix.lower()[1:]  # 점 제거
        if ext in self.supported_types:
            return ext
        
        # MIME 타입 기반
        mime_type, _ = mimetypes.guess_type(str(file_path))
        mime_to_ext = {
            'application/pdf': 'pdf',
            'text/plain': 'txt',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
            'text/csv': 'csv',
            'text/html': 'html',
            'text/markdown': 'md',
            'application/json': 'json'
        }
        
        return mime_to_ext.get(mime_type, 'txt')
    
    def _get_file_hash(self, file_path: Path) -> str:
        """파일 해시 계산"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    async def _load_by_type(self, file_path: Path, file_type: str) -> List[Document]:
        """파일 타입별 로딩"""
        loaders = {
            'pdf': self._load_pdf,
            'txt': self._load_text,
            'docx': self._load_docx,
            'xlsx': self._load_xlsx,
            'csv': self._load_csv,
            'html': self._load_html,
            'md': self._load_markdown,
            'json': self._load_json
        }
        
        loader = loaders.get(file_type)
        if not loader:
            raise ValueError(f"No loader available for file type: {file_type}")
        
        return await loader(file_path)
    
    async def _load_pdf(self, file_path: Path) -> List[Document]:
        """PDF 로딩"""
        documents = []
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={'page_number': page_num + 1}
                        ))
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
        
        return documents
    
    async def _load_text(self, file_path: Path) -> List[Document]:
        """텍스트 파일 로딩"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        return [Document(page_content=content, metadata={})]
    
    async def _load_docx(self, file_path: Path) -> List[Document]:
        """Word 문서 로딩"""
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n'.join(paragraphs)
        
        return [Document(page_content=content, metadata={})]
    
    async def _load_xlsx(self, file_path: Path) -> List[Document]:
        """Excel 파일 로딩"""
        documents = []
        
        # 모든 시트 읽기
        xl_file = pd.ExcelFile(file_path)
        
        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # 데이터프레임을 텍스트로 변환
            content_parts = []
            
            # 컬럼명 추가
            content_parts.append(f"시트: {sheet_name}")
            content_parts.append(f"컬럼: {', '.join(df.columns.tolist())}")
            
            # 각 행을 텍스트로 변환
            for idx, row in df.iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                
                if row_text:
                    content_parts.append(" | ".join(row_text))
            
            content = '\n'.join(content_parts)
            
            documents.append(Document(
                page_content=content,
                metadata={'sheet_name': sheet_name}
            ))
        
        return documents
    
    async def _load_csv(self, file_path: Path) -> List[Document]:
        """CSV 파일 로딩"""
        df = pd.read_csv(file_path)
        
        content_parts = []
        
        # 컬럼명 추가
        content_parts.append(f"컬럼: {', '.join(df.columns.tolist())}")
        
        # 각 행을 텍스트로 변환
        for idx, row in df.iterrows():
            row_text = []
            for col, value in row.items():
                if pd.notna(value):
                    row_text.append(f"{col}: {value}")
            
            if row_text:
                content_parts.append(" | ".join(row_text))
        
        content = '\n'.join(content_parts)
        
        return [Document(page_content=content, metadata={})]
    
    async def _load_html(self, file_path: Path) -> List[Document]:
        """HTML 파일 로딩"""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        # BeautifulSoup으로 텍스트 추출
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # 불필요한 태그 제거
        for script in soup(["script", "style"]):
            script.extract()
        
        text = soup.get_text()
        
        # 여러 공백을 하나로 정리
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return [Document(page_content=text, metadata={})]
    
    async def _load_markdown(self, file_path: Path) -> List[Document]:
        """Markdown 파일 로딩"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        # Markdown을 HTML로 변환한 후 텍스트 추출
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        return [Document(page_content=text, metadata={'format': 'markdown'})]
    
    async def _load_json(self, file_path: Path) -> List[Document]:
        """JSON 파일 로딩 (LangChain JSONLoader 사용)"""
        documents = []
        
        try:
            # JSON 파일 읽기
            with open(file_path, 'r', encoding='utf-8') as file:
                json_data = json.load(file)
            
            logger.info(f"JSON 파일 로드 시작: {file_path.name}")
            
            # JSON 데이터 타입별 처리
            if isinstance(json_data, list):
                # 리스트 형태의 JSON
                for idx, item in enumerate(json_data):
                    content = await self._process_json_item(item, idx)
                    if content:
                        documents.append(Document(
                            page_content=content,
                            metadata={
                                'json_type': 'list_item',
                                'item_index': idx,
                                'total_items': len(json_data)
                            }
                        ))
                        
            elif isinstance(json_data, dict):
                # 딕셔너리 형탄의 JSON
                content = await self._process_json_item(json_data, 0)
                if content:
                    documents.append(Document(
                        page_content=content,
                        metadata={
                            'json_type': 'object',
                            'keys': list(json_data.keys())[:10]  # 처음 10개 키만 저장
                        }
                    ))
                    
            else:
                # 기본 타입 (string, number 등)
                content = str(json_data)
                documents.append(Document(
                    page_content=content,
                    metadata={'json_type': 'primitive'}
                ))
            
            # 고급 처리: LangChain JSONLoader 사용 (선택적)
            if hasattr(self.document_config, 'json_loader_config'):
                json_config = self.document_config.json_loader_config
                if json_config.get('use_langchain_loader', False):
                    documents.extend(await self._load_json_with_langchain(file_path, json_config))
            
            logger.info(f"JSON 로딩 완료: {len(documents)}개 문서 생성")
            return documents
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON 파싱 오류: {e}")
            raise ValueError(f"Invalid JSON file: {e}")
        except Exception as e:
            logger.error(f"JSON 로딩 오류: {e}")
            raise
    
    async def _process_json_item(self, item: Any, index: int = 0) -> str:
        """개별 JSON 아이템을 텍스트로 변환"""
        try:
            if isinstance(item, dict):
                # 딕셔너리를 키-값 쌍으로 변환
                content_parts = []
                for key, value in item.items():
                    if isinstance(value, (dict, list)):
                        # 중첩된 객체는 JSON 문자열로 변환
                        value_str = json.dumps(value, ensure_ascii=False, indent=2)
                        content_parts.append(f"{key}: {value_str}")
                    else:
                        content_parts.append(f"{key}: {value}")
                return "\n".join(content_parts)
                
            elif isinstance(item, list):
                # 리스트를 인덱스가 있는 텍스트로 변환
                content_parts = []
                for idx, sub_item in enumerate(item):
                    if isinstance(sub_item, (dict, list)):
                        sub_content = json.dumps(sub_item, ensure_ascii=False, indent=2)
                    else:
                        sub_content = str(sub_item)
                    content_parts.append(f"[{idx}] {sub_content}")
                return "\n".join(content_parts)
                
            else:
                # 기본 타입
                return str(item)
                
        except Exception as e:
            logger.warning(f"JSON 아이템 처리 오류 (index {index}): {e}")
            return str(item)  # 오류 시 기본 문자열 변환
    
    async def _load_json_with_langchain(self, file_path: Path, json_config: Dict[str, Any]) -> List[Document]:
        """선택적 LangChain JSONLoader 사용 (고급 기능)"""
        try:
            # JSONLoader 설정
            jq_schema = json_config.get('jq_schema', '.')  # 전체 JSON 데이터
            content_key = json_config.get('content_key', None)  # 특정 키에서 콘텐츠 추출
            
            # LangChain JSONLoader 초기화
            if content_key:
                loader = JSONLoader(
                    file_path=str(file_path),
                    jq_schema=jq_schema,
                    content_key=content_key
                )
            else:
                loader = JSONLoader(
                    file_path=str(file_path),
                    jq_schema=jq_schema
                )
            
            # 문서 로드
            langchain_docs = await asyncio.to_thread(loader.load)
            
            # 메타데이터 추가
            for doc in langchain_docs:
                doc.metadata.update({
                    'json_loader': 'langchain',
                    'jq_schema': jq_schema,
                    'content_key': content_key
                })
            
            logger.info(f"LangChain JSONLoader로 {len(langchain_docs)}개 문서 추가 로드")
            return langchain_docs
            
        except Exception as e:
            logger.warning(f"LangChain JSONLoader 사용 실패, 기본 방식 사용: {e}")
            return []
    
    async def split_documents(self, documents: List[Document]) -> List[Document]:
        """문서 분할"""
        if not documents:
            return []
        
        logger.info(f"Splitting {len(documents)} documents into chunks")
        
        try:
            # RecursiveCharacterTextSplitter 사용
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=["\n\n", "\n", " ", ""]
            )
            
            # 분할 실행
            split_docs = splitter.split_documents(documents)
            
            # 청크 인덱스 추가
            for i, doc in enumerate(split_docs):
                doc.metadata['chunk_index'] = i
                doc.metadata['total_chunks'] = len(split_docs)
            
            logger.info(f"Documents split into {len(split_docs)} chunks")
            return split_docs
            
        except Exception as e:
            logger.error(f"Document splitting failed: {e}")
            raise
    
    async def embed_chunks(self, chunks: List[Document]) -> List[Dict[str, Any]]:
        """청크 임베딩 생성 (Dense + Sparse)"""
        if not chunks:
            return []
        
        logger.info(f"Generating dense and sparse embeddings for {len(chunks)} chunks")
        
        try:
            # 텍스트 추출
            texts = [chunk.page_content for chunk in chunks]
            
            # Dense 임베딩 생성 (배치 처리)
            dense_embeddings = await asyncio.to_thread(
                self.embedder.embed_documents, texts
            )
            logger.info(f"Dense embeddings generated: {len(dense_embeddings)} vectors")
            
            # Sparse 임베딩 생성 (BM42)
            sparse_embeddings = []
            if self.sparse_embedder:
                try:
                    sparse_results = await asyncio.to_thread(
                        list, self.sparse_embedder.embed(texts)
                    )
                    
                    # FastEmbed sparse 결과를 Qdrant 형식으로 변환
                    for sparse_result in sparse_results:
                        # sparse_result는 SparseEmbedding 객체
                        sparse_vector = {
                            "indices": sparse_result.indices.tolist(),
                            "values": sparse_result.values.tolist()
                        }
                        sparse_embeddings.append(sparse_vector)
                    
                    logger.info(f"Sparse embeddings generated: {len(sparse_embeddings)} BM42 vectors")
                    
                except Exception as e:
                    logger.warning(f"Sparse embedding generation failed, continuing with dense only: {e}")
                    sparse_embeddings = [None] * len(texts)
            else:
                logger.warning("Sparse embedder not available, using dense embeddings only")
                sparse_embeddings = [None] * len(texts)
            
            # 결과 구성
            embedded_chunks = []
            for chunk, dense_embedding, sparse_embedding in zip(chunks, dense_embeddings, sparse_embeddings):
                chunk_data = {
                    'content': chunk.page_content,
                    'dense_embedding': dense_embedding,
                    'metadata': chunk.metadata
                }
                
                # sparse embedding이 있는 경우에만 추가
                if sparse_embedding is not None:
                    chunk_data['sparse_embedding'] = sparse_embedding
                
                embedded_chunks.append(chunk_data)
            
            logger.info(f"Embeddings generated successfully for {len(embedded_chunks)} chunks")
            return embedded_chunks
            
        except Exception as e:
            logger.error(f"Embedding generation failed: {e}")
            raise
    
    async def process_document_full(self, file_path: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """문서 전체 처리 파이프라인"""
        try:
            # 1. 문서 로드
            documents = await self.load_document(file_path, metadata)
            
            # 2. 문서 분할
            chunks = await self.split_documents(documents)
            
            # 3. 임베딩 생성
            embedded_chunks = await self.embed_chunks(chunks)
            
            logger.info(f"Document processing completed: {file_path} -> {len(embedded_chunks)} embedded chunks")
            return embedded_chunks
            
        except Exception as e:
            logger.error(f"Full document processing failed for {file_path}: {e}")
            raise
    
    def get_stats(self) -> Dict[str, Any]:
        """통계 반환"""
        return {
            'supported_types': self.supported_types,
            'chunk_size': self.chunk_size,
            'chunk_overlap': self.chunk_overlap,
            'dense_embedder': {
                'model': self.embeddings_config.get('model', 'unknown'),
                'provider': self.embeddings_config.get('provider', 'unknown')
            },
            'sparse_embedder': {
                'model': self.embeddings_config.get('sparse_model', 'unknown'),
                'enabled': self.sparse_embedder is not None
            },
            'hybrid_search_enabled': self.sparse_embedder is not None
        }